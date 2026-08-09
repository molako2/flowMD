"""Export Word (.docx) : Markdown → pandoc, puis post-traitement RTL best-effort."""

from __future__ import annotations

import re
from pathlib import Path

_ARABIC_RE = re.compile(r"[؀-ۿݐ-ݿࢠ-ࣿﭐ-﷿ﹰ-﻿]")
_LATIN_RE = re.compile(r"[A-Za-zÀ-ɏ]")


def _is_arabic_dominant(text: str) -> bool:
    arabic = len(_ARABIC_RE.findall(text))
    latin = len(_LATIN_RE.findall(text))
    return arabic > 0 and arabic >= latin


def _apply_rtl(docx_path: Path) -> None:
    """Aligne à droite + active w:bidi pour les paragraphes majoritairement arabes."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = docx.Document(str(docx_path))

    def process(paragraphs) -> None:
        for para in paragraphs:
            if not _is_arabic_dominant(para.text):
                continue
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_pr = para._p.get_or_add_pPr()
            if p_pr.find(qn("w:bidi")) is None:
                bidi = p_pr.makeelement(qn("w:bidi"), {})
                p_pr.append(bidi)

    process(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process(cell.paragraphs)
    doc.save(str(docx_path))


def export_docx(md_path: Path, out_dir: Path, stem: str) -> tuple[Path, list[dict[str, str]]]:
    import pypandoc

    out_path = out_dir / f"{stem}.docx"
    warnings: list[dict[str, str]] = []

    pypandoc.convert_file(
        str(md_path),
        "docx",
        format="gfm",
        outputfile=str(out_path),
        extra_args=["--resource-path", str(out_dir)],
    )

    try:
        _apply_rtl(out_path)
    except Exception as exc:  # le RTL ne doit jamais faire échouer le job
        warnings.append(
            {
                "code": "DOCX_RTL_SKIPPED",
                "message": (
                    "Le document Word a été généré, mais la mise en forme "
                    f"droite-à-gauche (arabe) n'a pas pu être appliquée ({exc})."
                ),
            }
        )
    return out_path, warnings
