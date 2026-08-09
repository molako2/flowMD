"""Tests des exporteurs MD / DOCX / XLSX (sans docling, via FakeDocument)."""

import openpyxl
import pytest

from flowmd.exporters import export_all, normalize_formats
from flowmd.exporters.docx import _is_arabic_dominant
from flowmd.exporters.markdown import export_markdown
from flowmd.exporters.xlsx import export_xlsx


class TestNormalizeFormats:
    def test_aliases(self):
        assert normalize_formats("markdown,word,excel") == ["md", "docx", "xlsx"]

    def test_default_all(self):
        assert normalize_formats(None) == ["md", "docx", "xlsx"]
        assert normalize_formats("") == ["md", "docx", "xlsx"]

    def test_dedup(self):
        assert normalize_formats("md,md,docx") == ["md", "docx"]

    def test_unknown_raises(self):
        with pytest.raises(ValueError, match="inconnu"):
            normalize_formats("pdf")


class TestArabicDetection:
    def test_arabic_dominant(self):
        assert _is_arabic_dominant("تقرير مالي سنوي")

    def test_french_not_dominant(self):
        assert not _is_arabic_dominant("Facture des prestations")

    def test_mixed_mostly_arabic(self):
        assert _is_arabic_dominant("تقرير مالي سنوي 2026 MAD")

    def test_empty(self):
        assert not _is_arabic_dominant("")


class TestMarkdownExport:
    def test_fallback_export(self, fake_document, tmp_path):
        path = export_markdown(fake_document, tmp_path, "doc")
        assert path.name == "doc.md"
        content = path.read_text(encoding="utf-8")
        assert "FACTURE" in content
        assert "تقرير" in content


class TestXlsxExport:
    def test_tables_and_infos(self, fake_document, tmp_path):
        out = export_xlsx(
            fake_document,
            tmp_path / "doc.xlsx",
            {"source": "doc.pdf", "engine": "easyocr", "langs": ["fr"], "pages": 1},
        )
        wb = openpyxl.load_workbook(out)
        assert "Infos" in wb.sheetnames
        assert "Tableau_1" in wb.sheetnames
        sheet = wb["Tableau_1"]
        headers = [cell.value for cell in sheet[1]]
        assert "Désignation" in headers

    def test_no_tables(self, fake_document, tmp_path):
        fake_document.tables = []
        out = export_xlsx(fake_document, tmp_path / "vide.xlsx", {})
        wb = openpyxl.load_workbook(out)
        assert "Aucun tableau" in wb.sheetnames


class TestDocxExport:
    def test_full_export_with_rtl(self, fake_document, tmp_path):
        pytest.importorskip("pypandoc")
        outputs, warnings = export_all(
            fake_document, ["md", "docx"], tmp_path, "doc", {"source": "doc.pdf"}
        )
        assert outputs["docx"].is_file()
        assert not any(w["code"] == "DOCX_RTL_SKIPPED" for w in warnings)

        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        result = docx.Document(str(outputs["docx"]))
        arabic_paras = [p for p in result.paragraphs if "تقرير" in p.text]
        assert arabic_paras, "le paragraphe arabe doit exister dans le DOCX"
        assert arabic_paras[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


class TestExportAll:
    def test_md_only(self, fake_document, tmp_path):
        outputs, _ = export_all(fake_document, ["md"], tmp_path, "doc", {})
        assert set(outputs) == {"md"}

    def test_xlsx_only_still_writes_md_file(self, fake_document, tmp_path):
        # le markdown est produit en interne (aperçu) mais absent des sorties
        outputs, _ = export_all(fake_document, ["xlsx"], tmp_path, "doc", {})
        assert set(outputs) == {"xlsx"}
        assert (tmp_path / "doc.md").is_file()
